import streamlit as st
import pandas as pd
import numpy as np
import io
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

st.set_page_config(page_title="Weekly Update Report Tool", page_icon="📊", layout="wide")

st.title("📊 Weekly Update Report Tool")
st.markdown("Upload your three Excel files, fill in the details, and download your Word report + Excel file — no coding needed.")
st.divider()

# ── Helper: init session state ────────────────────────────────────────────────
def ss(key, default):
    if key not in st.session_state:
        st.session_state[key] = default

ss("num_sites", 3)
ss("pipeline_done", False)
ss("site_tables", {})
ss("activity_tags", {})
ss("excel_bytes", None)
ss("word_bytes", None)
ss("df_totals", None)
ss("result", None)
ss("missing_summary", None)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — Upload files
# ══════════════════════════════════════════════════════════════════════════════
st.subheader("Step 1 · Upload Your Excel Files")
col1, col2, col3 = st.columns(3)
with col1:
    students_file = st.file_uploader("🎒 Students File", type=["xlsx","xls"], key="students")
with col2:
    adults_file   = st.file_uploader("👪 Adults / Family File", type=["xlsx","xls"], key="adults")
with col3:
    all_file      = st.file_uploader("📋 All Sites File", type=["xlsx","xls"], key="all")

st.divider()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — Program info + targets
# ══════════════════════════════════════════════════════════════════════════════
st.subheader("Step 2 · Program Info & Targets")

program_name = st.text_input("Program / District Name (for report header)",
                              placeholder="e.g. SY2026 - CSD 3 21st CCLC Program")
report_date  = st.text_input("Report Date", placeholder="e.g. March 16, 2026")

st.markdown("#### Per-Site Targets")
st.info("Enter one row per site — in the same order as they appear in your Students file. Do not include the Total row.")

col_add, col_rem, _ = st.columns([1,1,4])
with col_add:
    if st.button("➕ Add site"):
        st.session_state.num_sites += 1
with col_rem:
    if st.button("➖ Remove last") and st.session_state.num_sites > 1:
        st.session_state.num_sites -= 1

site_rows = []
defaults_students = [150, 200, 100]
defaults_parents  = [24, 70, 50]

h1, h2, h3, h4 = st.columns([2, 1, 1, 1])
h1.markdown("**Site name** *(optional label)*")
h2.markdown("**Target # Students**")
h3.markdown("**Target # Parents**")
h4.markdown("**Target Literacy Workshops**")

for i in range(st.session_state.num_sites):
    c1, c2, c3, c4 = st.columns([2, 1, 1, 1])
    name = c1.text_input(f"site_name_{i}", key=f"sname_{i}",
                          label_visibility="collapsed",
                          placeholder=f"Site {i+1}")
    t_students = c2.number_input(f"ts_{i}", min_value=0,
                                  value=defaults_students[i] if i < 3 else 100,
                                  label_visibility="collapsed", key=f"ts_{i}")
    t_parents  = c3.number_input(f"tp_{i}", min_value=0,
                                  value=defaults_parents[i] if i < 3 else 50,
                                  label_visibility="collapsed", key=f"tp_{i}")
    t_lit      = c4.number_input(f"tl_{i}", min_value=0, value=0,
                                  label_visibility="collapsed", key=f"tl_{i}")
    site_rows.append({
        "name": name or f"Site {i+1}",
        "target_students": int(t_students),
        "target_parents":  int(t_parents),
        "target_lit":      int(t_lit),
    })

st.divider()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — Run pipeline
# ══════════════════════════════════════════════════════════════════════════════
st.subheader("Step 3 · Process Files")
all_uploaded = students_file and adults_file and all_file

if st.button("⚙️ Process Files", disabled=not all_uploaded, use_container_width=True):
    st.session_state.pipeline_done = False
    st.session_state.activity_tags = {}
    st.session_state.word_bytes = None

    with st.status("Running pipeline…", expanded=True) as status:
        try:
            target_values = [r["target_students"] for r in site_rows]

            st.write("📂 Reading files…")
            students_bytes = students_file.read()
            adults_bytes   = adults_file.read()
            all_bytes      = all_file.read()
            students_sheets = pd.read_excel(io.BytesIO(students_bytes), sheet_name=None)

            # ── Student Summary ───────────────────────────────────────────────
            st.write("🍀 Student Summary Statistics…")
            df_pbh = students_sheets['Participants By Hour Band']
            df_pbh.columns = df_pbh.iloc[4]
            df_pbh = df_pbh.iloc[5:].reset_index(drop=True)

            dsa = students_sheets['Daily Site Attendance Summary']
            dsa.columns = dsa.iloc[2]
            dsa = dsa.iloc[3:]
            dsa.columns.name = None
            dsa = dsa.reset_index(drop=True)[['Total']].iloc[:-1]
            dsa['Total'] = dsa['Total'].str.extract(r'(\d+\.?\d*)')

            all_cols    = ['0','Less Than 15','15-44','45-89','90-179','180-269','270+']
            served_cols = ['Less Than 15','15-44','45-89','90-179','180-269','270+']
            plus15_cols = ['15-44','45-89','90-179','180-269','270+']
            plus90_cols = ['90-179','180-269','270+']
            ec = df_pbh.columns.tolist()
            all_cols    = [c for c in all_cols    if c in ec]
            served_cols = [c for c in served_cols if c in ec]
            plus15_cols = [c for c in plus15_cols if c in ec]
            plus90_cols = [c for c in plus90_cols if c in ec]
            df_pbh[all_cols] = df_pbh[all_cols].apply(pd.to_numeric, errors='coerce').fillna(0)
            df_sub = df_pbh[df_pbh['Site'] == 'Subtotal']

            df_totals = pd.DataFrame({
                '[Total # Enrolled]': df_sub[all_cols].sum(axis=1),
                '[Total # Served]':   df_sub[served_cols].sum(axis=1),
                '[Total # 15+]':      df_sub[plus15_cols].sum(axis=1),
                '[Total # 90+]':      df_sub[plus90_cols].sum(axis=1),
            })
            df_totals.insert(0, '[Target # of students served]', target_values)
            dsa['Total'] = dsa['Total'].astype(int)
            dsa = dsa.rename(columns={'Total': 'Avg. # of Students Per Day'})
            df_totals.insert(2, 'Avg. # of Students Per Day', dsa['Avg. # of Students Per Day'].values)

            school_names = [
                row['Site'] for _, row in df_pbh.iterrows()
                if pd.notna(row['Site']) and row['Site'] != 'Subtotal' and row['Site'] != 'Total'
            ]
            df_totals.insert(0, 'School', school_names)
            total_row = pd.DataFrame(df_totals.iloc[:, 1:].sum()).T
            total_row.insert(0, 'School', 'Total')
            df_totals = pd.concat([df_totals, total_row], ignore_index=True)
            df_totals['# of students 15+ hrs total (% of Target)'] = (
                df_totals['[Total # 15+]'].astype(int).astype(str) + " (" +
                ((df_totals['[Total # 15+]'] / df_totals['[Target # of students served]']) * 100)
                .round().astype(int).astype(str) + "%)"
            )
            df_totals['# of students 90+ hrs total (% of Target)'] = (
                df_totals['[Total # 90+]'].astype(int).astype(str) + " (" +
                ((df_totals['[Total # 90+]'] / df_totals['[Target # of students served]']) * 100)
                .round().astype(int).astype(str) + "%)"
            )
            df_totals = df_totals.drop(columns=['[Total # 15+]','[Total # 90+]']).reset_index(drop=True)

            # ── Family Component ──────────────────────────────────────────────
            st.write("🌷 Family Component…")
            df_hours = pd.read_excel(io.BytesIO(adults_bytes), sheet_name='Participant Attendance Hours', skiprows=2)
            df_hours['HoursPresent'] = pd.to_numeric(df_hours['HoursPresent'], errors='coerce')
            df_hours['ParticipantId'] = df_hours['ParticipantId'].astype(str).str.replace(r'\.0$','',regex=True)
            df_active = df_hours[(df_hours['HoursPresent'] > 0) & (df_hours['ParticipantId'].str.len() != 9)]
            result = df_active.groupby('Site')['ParticipantId'].nunique().reset_index()
            result.rename(columns={'ParticipantId': 'Parents Served (Total)'}, inplace=True)
            result.loc[len(result)] = {'Site':'Total','Parents Served (Total)': result['Parents Served (Total)'].sum()}

            # ── Demographics / Missing ────────────────────────────────────────
            st.write("🌸 Participant Demographics…")
            df_demo = students_sheets['Participant Demographics']
            df_demo.columns = df_demo.iloc[2]
            df_demo = df_demo.iloc[3:].reset_index(drop=True)

            def summarize_missing(df, cols_check, cat='Site'):
                missing_site_rows = df[df[cat].isna() | (df[cat].astype(str).str.strip() == '')].copy()
                sub = df[cols_check + [cat]].copy()
                sfg = sub[sub[cat].notna()].copy()
                sfg[cat] = sfg[cat].astype(str).str.title()
                for col in cols_check:
                    cl = sfg[col].astype(str).str.strip()
                    nem = cl.str.lower() == 'not entered'
                    if col == 'Gender':
                        vg = cl.str.title().isin(['Male','Female','Non-Binary'])
                        sfg[col+'_missing'] = ((~vg)|nem).astype(int)
                    else:
                        sfg[col+'_missing'] = (sfg[col].isna()|nem).astype(int)
                pid  = df.loc[sfg.index,'ParticipantID'].astype(str).str.strip()
                spid = df.loc[sfg.index,'State ParticipantID'].astype(str).str.strip()
                vp = pid.str.match(r'^[12]\d{8}$')
                vs = spid.str.match(r'^[12]\d{8}$')
                sfg['ParticipantID_missing']       = (~vp).astype(int)
                sfg['State ParticipantID_missing'] = (~vs).astype(int)
                sfg['OSIS_missing']                = (~((pid==spid)&vp&vs)).astype(int)
                mc = [c+'_missing' for c in cols_check] + ['ParticipantID_missing','State ParticipantID_missing','OSIS_missing']
                pivot = sfg.groupby(cat)[mc].sum().reset_index()
                tr = pd.DataFrame(pivot[mc].sum()).T
                tr[cat] = 'Total'
                pivot = pd.concat([pivot, tr], ignore_index=True)
                amf = df.copy()
                pa  = amf['ParticipantID'].astype(str).str.strip()
                spa = amf['State ParticipantID'].astype(str).str.strip()
                vpa = pa.str.match(r'^[12]\d{8}$')
                vsa = spa.str.match(r'^[12]\d{8}$')
                for col in cols_check:
                    cl = amf[col].astype(str).str.strip()
                    nem = cl.str.lower() == 'not entered'
                    if col == 'Gender':
                        vg = cl.str.title().isin(['Male','Female','Non-Binary'])
                        amf[col+'_missing'] = ((~vg)|nem).astype(int)
                    else:
                        amf[col+'_missing'] = (amf[col].isna()|nem).astype(int)
                amf['ParticipantID_missing']       = (~vpa).astype(int)
                amf['State ParticipantID_missing'] = (~vsa).astype(int)
                amf['OSIS_missing']                = (~((pa==spa)&vpa&vsa)).astype(int)
                dob_p = pd.to_datetime(amf['Date Of Birth'], errors='coerce')
                amf['DOB_too_young'] = (dob_p.dt.year > 2023).astype(int)
                fc2 = [c+'_missing' for c in cols_check] + ['ParticipantID_missing']
                tmr = amf[amf[fc2].sum(axis=1) > 0].copy()
                ydr = amf[amf['DOB_too_young'] == 1].copy()
                return pivot, missing_site_rows, tmr, fc2, ydr

            missing_summary, missing_site_rows, total_missing_rows, flag_cols, young_dob_rows = \
                summarize_missing(df_demo, ['Date Of Birth','Grade Level','Gender'])

            # ── Site Summary ──────────────────────────────────────────────────
            st.write("🪻 Site Summary Report…")
            all_io = io.BytesIO(all_bytes)
            df_act = pd.read_excel(all_io, sheet_name='Activity-Session Details', skiprows=2)
            all_io.seek(0)
            df_enr = pd.read_excel(all_io, sheet_name='Session Enrollment by Session', skiprows=2)
            all_io.seek(0)
            df_att = pd.read_excel(all_io, sheet_name='Daily Activity Attendance Summa', skiprows=4)

            df_act = df_act[['Site','Activity','Session','Days Scheduled','Session Start Date']].copy()
            df_act['Days Scheduled'] = pd.to_numeric(df_act['Days Scheduled'], errors='coerce')
            df_act['Session Start Date'] = pd.to_datetime(df_act['Session Start Date'], errors='coerce')
            today = pd.Timestamp.today().normalize()
            df_enr = df_enr[['Site','Activity','Session','Enrolled Count']].copy()
            df_enr['Enrolled Count'] = pd.to_numeric(df_enr['Enrolled Count'], errors='coerce')
            df_enr.rename(columns={'Enrolled Count':'Enrolled Participant'}, inplace=True)

            def extract_avg(v):
                if pd.isna(v): return np.nan
                try: return float(str(v).replace('Average:','').strip())
                except: return np.nan

            df_att = df_att[['Site','Activity','Session','Total']].copy()
            df_att['Total'] = df_att['Total'].apply(extract_avg)
            df_att.rename(columns={'Total':'Average Daily Attendance'}, inplace=True)

            sites = [s for s in df_act['Site'].dropna().unique()
                     if str(s).strip() != '' and not str(s).startswith('Total') and not str(s).startswith('Average')]
            site_tables = {}
            for site in sites:
                m = pd.merge(df_act[df_act['Site']==site], df_enr[df_enr['Site']==site],
                             on=['Site','Activity','Session'], how='outer')
                m = pd.merge(m, df_att[df_att['Site']==site],
                             on=['Site','Activity','Session'], how='outer')
                m = m[~(m['Session Start Date'] >= today)].drop(columns=['Session Start Date'], errors='ignore')
                m = m.fillna('-').sort_values(['Activity','Session']).reset_index(drop=True)
                site_tables[site] = m

            # ── Write Excel ───────────────────────────────────────────────────
            st.write("🌈 Writing Excel…")
            out_buf = io.BytesIO()
            with pd.ExcelWriter(out_buf, engine='openpyxl') as writer:
                df_totals.to_excel(writer, sheet_name='Student Statistics', index=False)
                result.to_excel(writer, sheet_name='Parents Served Summary', index=False)
                missing_summary.to_excel(writer, sheet_name='Missing Summary', index=False)
                missing_site_rows.to_excel(writer, sheet_name='Missing Site Info', index=False)
                dc = [c for c in total_missing_rows.columns if not c.endswith('_missing') and c != 'DOB_too_young']
                total_missing_rows[dc].to_excel(writer, sheet_name='Total Missing Info', index=False)
                ydc = [c for c in young_dob_rows.columns if not c.endswith('_missing') and c != 'DOB_too_young']
                young_dob_rows[ydc].to_excel(writer, sheet_name='Young DOB', index=False)
                for sn, fd in site_tables.items():
                    safe = str(sn)[:31].replace(':','').replace('/','').replace('\\','').replace('?','').replace('*','')
                    fd.to_excel(writer, sheet_name=safe, index=False)

            out_buf.seek(0)
            wb = load_workbook(out_buf)
            red_fill  = PatternFill('solid', start_color='FF9999', end_color='FF9999')
            blue_fill = PatternFill('solid', start_color='9999FF', end_color='9999FF')
            f2o = {fc: fc[:-len('_missing')] for fc in flag_cols if fc.endswith('_missing')}
            ws = wb['Total Missing Info']
            hdr = {cell.value: cell.column for cell in ws[1]}
            for ridx, (_, row) in enumerate(total_missing_rows.iterrows(), start=2):
                for fc, oc in f2o.items():
                    if oc in hdr and row.get(fc, 0) == 1:
                        ws.cell(row=ridx, column=hdr[oc]).fill = red_fill
            ws2 = wb['Young DOB']
            hdr2 = {cell.value: cell.column for cell in ws2[1]}
            if 'Date Of Birth' in hdr2:
                for ridx in range(2, len(young_dob_rows)+2):
                    ws2.cell(row=ridx, column=hdr2['Date Of Birth']).fill = blue_fill
            final_buf = io.BytesIO()
            wb.save(final_buf)

            # ── Save to session state ─────────────────────────────────────────
            st.session_state.excel_bytes     = final_buf.getvalue()
            st.session_state.df_totals       = df_totals
            st.session_state.result          = result
            st.session_state.missing_summary = missing_summary
            st.session_state.site_tables     = site_tables

            for site, df in site_tables.items():
                acts = sorted(df['Activity'].replace('-', np.nan).dropna().unique().tolist())
                if site not in st.session_state.activity_tags:
                    st.session_state.activity_tags[site] = {a: "Neither" for a in acts}

            st.session_state.pipeline_done = True
            status.update(label="✅ Files processed!", state="complete")

        except Exception as e:
            status.update(label="❌ Error", state="error")
            st.error(f"Error: {e}")
            st.exception(e)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — Tag activities
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.pipeline_done:
    st.divider()
    st.subheader("Step 4 · Tag Family Activities")
    st.info("For each site, mark which activities count as **Literacy Workshops** or **Family Engagement Events**. Everything else stays as 'Neither'.")

    tags = st.session_state.activity_tags
    for site, act_dict in tags.items():
        with st.expander(f"📍 {site}", expanded=False):
            acts = list(act_dict.keys())
            ch1, ch2 = st.columns([3, 2])
            ch1.markdown("**Activity**")
            ch2.markdown("**Type**")
            for act in acts:
                ac1, ac2 = st.columns([3, 2])
                ac1.markdown(f"<small>{act}</small>", unsafe_allow_html=True)
                choice = ac2.selectbox(
                    label=act,
                    options=["Neither","Literacy Workshop","Family Engagement Event"],
                    index=["Neither","Literacy Workshop","Family Engagement Event"].index(act_dict.get(act,"Neither")),
                    key=f"tag_{site}_{act}",
                    label_visibility="collapsed"
                )
                tags[site][act] = choice

    st.divider()

    # ══════════════════════════════════════════════════════════════════════════
    # STEP 5 — Generate Word + download
    # ══════════════════════════════════════════════════════════════════════════
    st.subheader("Step 5 · Generate Report")

    if st.button("📄 Generate Word Doc + Excel", type="primary", use_container_width=True):
        with st.spinner("Building Word document…"):
            try:
                from docx import Document as DocxDoc
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement

                df_totals    = st.session_state.df_totals
                result       = st.session_state.result
                missing_sum  = st.session_state.missing_summary
                site_tables  = st.session_state.site_tables
                tags         = st.session_state.activity_tags

                # ── Build family rows ─────────────────────────────────────────
                family_rows = []
                result_no_total = result[result['Site'] != 'Total'].reset_index(drop=True)

                for i, sr in enumerate(site_rows):
                    site_label = sr["name"]
                    matched = None
                    for k in site_tables.keys():
                        if site_label.lower() in k.lower() or k.lower() in site_label.lower():
                            matched = k; break
                    if matched is None and i < len(list(site_tables.keys())):
                        matched = list(site_tables.keys())[i]

                    workshops = engagement = 0
                    if matched and matched in tags:
                        sdf = site_tables[matched]
                        for act, tag in tags[matched].items():
                            act_rows = sdf[sdf['Activity'] == act]
                            cnt = len(act_rows[act_rows['Days Scheduled'] != '-'])
                            if tag == "Literacy Workshop":      workshops  += cnt
                            elif tag == "Family Engagement Event": engagement += cnt

                    parents_served = int(result_no_total.iloc[i]['Parents Served (Total)']) if i < len(result_no_total) else 0
                    t_par = sr["target_parents"]
                    pct = f"{round((parents_served/t_par)*100)}%" if t_par > 0 else "N/A"
                    family_rows.append({
                        "school": site_label,
                        "target_lit": sr["target_lit"] if sr["target_lit"] > 0 else "N/A",
                        "completed_workshops": workshops,
                        "engagement_events": engagement,
                        "target_parents": t_par,
                        "parents_served": f"{parents_served} ({pct})",
                    })

                fam_tot_ws  = sum(r["completed_workshops"] for r in family_rows)
                fam_tot_eng = sum(r["engagement_events"]   for r in family_rows)
                fam_tot_tp  = sum(r["target_parents"]      for r in family_rows)
                fam_tot_ps  = int(result[result['Site']=='Total']['Parents Served (Total)'].values[0])
                fam_tot_pct = f"{round((fam_tot_ps/fam_tot_tp)*100)}%" if fam_tot_tp > 0 else "N/A"

                # ── Word helpers ──────────────────────────────────────────────
                def set_cell_bg(cell, hex_color):
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), hex_color)
                    tcPr.append(shd)

                def set_borders(cell):
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcB = OxmlElement('w:tcBorders')
                    for side in ['top','left','bottom','right']:
                        b = OxmlElement(f'w:{side}')
                        b.set(qn('w:val'), 'single')
                        b.set(qn('w:sz'), '4')
                        b.set(qn('w:space'), '0')
                        b.set(qn('w:color'), 'CCCCCC')
                        tcB.append(b)
                    tcPr.append(tcB)

                def hcell(cell, text, fs=9, bg='1F4E79', fg='FFFFFF'):
                    set_cell_bg(cell, bg)
                    set_borders(cell)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(text)
                    r.bold = True
                    r.font.size = Pt(fs)
                    r.font.color.rgb = RGBColor.from_string(fg)

                def dcell(cell, text, bold=False, fs=9, center=True, bg=None):
                    if bg: set_cell_bg(cell, bg)
                    set_borders(cell)
                    p = cell.paragraphs[0]
                    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(str(text))
                    r.bold = bold
                    r.font.size = Pt(fs)

                def add_section_heading(doc, text):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(12)
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                    return p

                # ── Build doc ─────────────────────────────────────────────────
                doc = DocxDoc()
                for section in doc.sections:
                    section.top_margin    = Inches(0.75)
                    section.bottom_margin = Inches(0.75)
                    section.left_margin   = Inches(0.75)
                    section.right_margin  = Inches(0.75)

                # Title block
                for text, size, bold in [
                    (program_name or "Weekly Update Report", 14, True),
                    ("Weekly AfterSchool21 Data Update", 12, True),
                    (report_date or "", 10, False),
                ]:
                    if text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run(text)
                        r.bold = bold
                        r.font.size = Pt(size)
                        if bold: r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

                doc.add_paragraph()

                # ── 1. Student Summary ────────────────────────────────────────
                add_section_heading(doc, "Student Summary Statistics")
                stu_cols = ['School','[Target # of students served]','[Total # Enrolled]',
                            '[Total # Served]','Avg. # of Students Per Day',
                            '# of students 15+ hrs total (% of Target)',
                            '# of students 90+ hrs total (% of Target)']
                stu_hdrs = ['School','Target #\nStudents','Total #\nEnrolled','Total #\nServed',
                            'Avg. #\nPer Day','# Students 15+\nhrs (% Target)','# Students 90+\nhrs (% Target)']
                stu_ws   = [1.9, 0.75, 0.75, 0.75, 0.75, 1.1, 1.1]
                stu_df   = df_totals[stu_cols]

                t = doc.add_table(rows=1+len(stu_df), cols=len(stu_cols))
                t.style = 'Table Grid'
                for j, (h, w) in enumerate(zip(stu_hdrs, stu_ws)):
                    hcell(t.rows[0].cells[j], h)
                    t.columns[j].width = Inches(w)
                for i, (_, row) in enumerate(stu_df.iterrows(), start=1):
                    is_tot = str(row['School']) == 'Total'
                    bg = 'D6E4F0' if is_tot else None
                    for j, col in enumerate(stu_cols):
                        dcell(t.rows[i].cells[j], row[col], bold=is_tot, bg=bg, center=(j>0))

                doc.add_paragraph()

                # ── 2. Family Component ───────────────────────────────────────
                add_section_heading(doc, "Family Component Summary Statistics")
                fam_hdrs = ['School – CBO Provider','Target\nLiteracy\nWorkshops',
                            'Completed\nWorkshops','Family\nEngagement\nEvents',
                            'Target #\nParents','Parents Served\n(% of Target)']
                fam_ws   = [2.1, 0.85, 0.85, 0.85, 0.75, 1.2]

                ft = doc.add_table(rows=1+len(family_rows)+1, cols=6)
                ft.style = 'Table Grid'
                for j, (h, w) in enumerate(zip(fam_hdrs, fam_ws)):
                    hcell(ft.rows[0].cells[j], h)
                    ft.columns[j].width = Inches(w)
                for i, fr in enumerate(family_rows, start=1):
                    dcell(ft.rows[i].cells[0], fr["school"], center=False)
                    dcell(ft.rows[i].cells[1], fr["target_lit"])
                    dcell(ft.rows[i].cells[2], fr["completed_workshops"])
                    dcell(ft.rows[i].cells[3], fr["engagement_events"])
                    dcell(ft.rows[i].cells[4], fr["target_parents"])
                    dcell(ft.rows[i].cells[5], fr["parents_served"])
                tri = len(family_rows)+1
                for j, v in enumerate(['Total','',fam_tot_ws,fam_tot_eng,fam_tot_tp,
                                        f"{fam_tot_ps} ({fam_tot_pct})"]):
                    dcell(ft.rows[tri].cells[j], v, bold=True, bg='D6E4F0', center=(j>0))

                doc.add_paragraph()

                # ── 3. Missing Info ───────────────────────────────────────────
                add_section_heading(doc, "Missing Student Information")
                ms = missing_sum.copy().rename(columns={
                    'Site':'School','Date Of Birth_missing':'Missing DOB',
                    'Grade Level_missing':'Missing Grade','Gender_missing':'Missing Gender',
                    'ParticipantID_missing':'Missing OSIS #',
                    'State ParticipantID_missing':'Missing State ID',
                    'OSIS_missing':'OSIS Mismatch',
                })
                ms_cols = ['School','Missing DOB','Missing Grade','Missing Gender',
                           'Missing OSIS #','Missing State ID','OSIS Mismatch']
                ms_cols = [c for c in ms_cols if c in ms.columns]
                ms_ws   = [2.0] + [0.85]*(len(ms_cols)-1)

                mt = doc.add_table(rows=1+len(ms), cols=len(ms_cols))
                mt.style = 'Table Grid'
                for j, (h, w) in enumerate(zip(ms_cols, ms_ws)):
                    hcell(mt.rows[0].cells[j], h)
                    mt.columns[j].width = Inches(w)
                for i, (_, row) in enumerate(ms.iterrows(), start=1):
                    is_tot = str(row.get('School','')) == 'Total'
                    bg = 'D6E4F0' if is_tot else None
                    for j, col in enumerate(ms_cols):
                        dcell(mt.rows[i].cells[j], row[col], bold=is_tot, bg=bg, center=(j>0))

                doc.add_paragraph()

                # ── 4. Site Summary Reports ───────────────────────────────────
                add_section_heading(doc, "Site Summary Reports: Attendance by Activity")

                for site, sdf in site_tables.items():
                    p = doc.add_paragraph()
                    r = p.add_run(site)
                    r.bold = True
                    r.font.size = Pt(11)
                    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                    note = doc.add_paragraph()
                    note.add_run("*Family Literacy and Engagement Workshops are highlighted in blue.").italic = True
                    note.runs[0].font.size = Pt(8)
                    note.runs[0].font.color.rgb = RGBColor(0x26, 0x4F, 0x78)

                    site_cols = [c for c in ['Activity','Session','Days Scheduled',
                                              'Enrolled Participant','Average Daily Attendance']
                                 if c in sdf.columns]
                    site_hdr_map = {
                        'Activity':'Activity','Session':'Session',
                        'Days Scheduled':'# Days\nScheduled',
                        'Enrolled Participant':'# Enrolled\nParticipant',
                        'Average Daily Attendance':'Avg Daily\nAttendance'
                    }
                    site_w_map = {'Activity':2.2,'Session':2.2,'Days Scheduled':0.8,
                                  'Enrolled Participant':0.9,'Average Daily Attendance':0.9}

                    st2 = doc.add_table(rows=1+len(sdf), cols=len(site_cols))
                    st2.style = 'Table Grid'
                    for j, col in enumerate(site_cols):
                        hcell(st2.rows[0].cells[j], site_hdr_map.get(col,col), fs=8)
                        st2.columns[j].width = Inches(site_w_map.get(col,1.0))

                    site_tag_dict = tags.get(site, {})
                    for i, (_, row) in enumerate(sdf.iterrows(), start=1):
                        act = row.get('Activity','-')
                        tag = site_tag_dict.get(act,'Neither')
                        rbg = 'BDD7EE' if tag in ['Literacy Workshop','Family Engagement Event'] else None
                        for j, col in enumerate(site_cols):
                            dcell(st2.rows[i].cells[j], row[col], fs=8, bg=rbg, center=(j>1))

                    doc.add_paragraph()

                word_buf = io.BytesIO()
                doc.save(word_buf)
                st.session_state.word_bytes = word_buf.getvalue()
                st.success("🎉 Both files are ready to download!")

            except Exception as e:
                st.error(f"Error generating Word doc: {e}")
                st.exception(e)

    # ── Download buttons ──────────────────────────────────────────────────────
    if st.session_state.word_bytes or st.session_state.excel_bytes:
        st.divider()
        dl1, dl2 = st.columns(2)
        with dl1:
            if st.session_state.word_bytes:
                st.download_button(
                    "⬇️ Download Word Report (.docx)",
                    data=st.session_state.word_bytes,
                    file_name="Weekly_Update_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary"
                )
        with dl2:
            if st.session_state.excel_bytes:
                st.download_button(
                    "⬇️ Download Excel Data (.xlsx)",
                    data=st.session_state.excel_bytes,
                    file_name="Weekly_Update_Report.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )
